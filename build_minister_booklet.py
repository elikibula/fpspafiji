from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "booklet"
OUT.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / "static/assets/images/fpspa_logo.png"
IMAGES = ROOT / "static/assets/images/fpspa"

NAVY = "0F2F5F"
DEEP = "06162F"
BLUE = "1F5CA8"
GOLD = "FDE339"
PALE = "EEF3F8"
INK = "24364B"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(p, text, size=10, color=INK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return r


def kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, text.upper(), 8, BLUE, True)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def body(doc, text, lead=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    add_text(p, text, 11 if lead else 9.5, INK, lead)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    add_text(p, text, 9.3, INK)


def add_photo(doc, filename, width=4.55):
    image_path = IMAGES / filename
    if image_path.suffix.lower() == ".webp":
        converted = OUT / (image_path.stem + ".png")
        if not converted.exists():
            Image.open(image_path).convert("RGB").save(converted, "PNG")
        image_path = converted
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    p.add_run().add_picture(str(image_path), width=Inches(width))


def stats(doc):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    values = [("735", "Heads of Schools"), ("12", "District networks"), ("43", "Years established")]
    for cell, (number, label) in zip(table.rows[0].cells, values):
        cell.width = Inches(1.55)
        shade(cell, NAVY)
        margins(cell, 150, 80, 150, 80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, number, 19, GOLD, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p2, label, 7.5, WHITE, True)


def page_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "FIJI HEAD TEACHERS ASSOCIATION  |  LEADING SCHOOLS. SHAPING FUTURES.", 7, BLUE, True)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(5.83)
sec.page_height = Inches(8.27)
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.62)
sec.right_margin = Inches(0.62)
sec.header_distance = Inches(0.25)
sec.footer_distance = Inches(0.25)
page_footer(sec)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12
for name, size, before, after in (("Heading 1", 19, 8, 7), ("Heading 2", 13, 7, 4)):
    st = styles[name]
    st.font.name = "Aptos Display"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(NAVY)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
lb = styles["List Bullet"]
lb.font.name = "Aptos"
lb.font.size = Pt(9.3)
lb.paragraph_format.left_indent = Inches(0.24)
lb.paragraph_format.first_line_indent = Inches(-0.14)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.add_run().add_picture(str(LOGO), width=Inches(1.35))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
add_text(p, "FIJI HEAD TEACHERS ASSOCIATION", 10, BLUE, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(7)
add_text(p, "Strengthening Primary\nSchool Leadership", 25, NAVY, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "A briefing booklet for the Minister for Education", 12, INK, True)
add_photo(doc, "conference-leaders.jpg", 4.5)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "Leadership  •  Advocacy  •  Excellence", 10, BLUE, True)

# Page 2
doc.add_page_break()
kicker(doc, "Our national leadership network")
heading(doc, "One association. A nationwide voice.")
body(doc, "The Fiji Head Teachers Association (FHTA) connects primary school leaders across Fiji, representing their professional interests and strengthening the leadership that shapes schools, teachers and student outcomes.", True)
stats(doc)
heading(doc, "What FHTA contributes", 2)
bullet(doc, "A trusted channel between Heads of Schools, communities, partners and the Ministry of Education.")
bullet(doc, "Professional learning and peer networks that translate policy into effective school-level practice.")
bullet(doc, "National, divisional and district representation on the issues affecting primary education.")
bullet(doc, "Practical member support, shared resources and timely communication.")
add_photo(doc, "principals-audience.webp", 4.55)

# Page 3
doc.add_page_break()
kicker(doc, "Purpose in practice")
heading(doc, "Leading schools. Shaping futures.")
body(doc, "FHTA's work is grounded in the daily realities of school leadership. The Association helps Heads of Schools lead teaching and learning, manage people and resources, respond to community needs, and sustain improvement.", True)
heading(doc, "Six areas of service", 2)
for title, detail in [
    ("Advocacy and representation", "A coordinated voice at district, divisional and national levels."),
    ("Professional concerns", "Support pathways for members facing workplace and leadership issues."),
    ("Professional learning", "Workshops, online learning and capacity building."),
    ("Leadership practice", "Approaches that strengthen teaching, learning and school operations."),
    ("Collaboration", "Peer exchange through area, branch and national networks."),
    ("Ministry partnership", "Constructive engagement in education initiatives and policy implementation."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, title + ". ", 9.4, NAVY, True)
    add_text(p, detail, 9.4, INK)

# Page 4
doc.add_page_break()
kicker(doc, "Digital association platform")
heading(doc, "A connected website for public information and member service")
body(doc, "The FHTA website combines a public communications channel with secure member and staff services. It gives the Association one place to inform, support, train and connect school leaders.", True)
add_photo(doc, "homepage.png", 4.5)
heading(doc, "Three connected experiences", 2)
for title, detail in [
    ("Public website", "Association information, news, events, services, resources, leadership and district coverage."),
    ("Member portal", "Profiles, membership status, documents, helpdesk support, representatives and learning."),
    ("Staff operations", "Member approvals, content publishing, events, resources, support tickets, training, attendance and certificates."),
]:
    p = doc.add_paragraph()
    add_text(p, title + " — ", 9.3, BLUE, True)
    add_text(p, detail, 9.3, INK)

# Page 5
doc.add_page_break()
kicker(doc, "Professional learning")
heading(doc, "From enrolment to evidence of completion")
body(doc, "The training platform supports online, workshop and blended learning. Members can move through a clear pathway while staff manage delivery and participation.", True)
add_photo(doc, "district-symposium.webp", 4.55)
for n, text in enumerate([
    "Browse published courses and select an available schedule.",
    "Enrol, access lessons, resources and learning outcomes.",
    "Complete quizzes and track progress.",
    "Check in at workshops using authenticated QR attendance.",
    "Receive a printable certificate when completion requirements are met.",
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"{n:02d}", 11, BLUE, True)
    add_text(p, "   " + text, 9.4, INK)

# Page 6
doc.add_page_break()
kicker(doc, "School leadership priorities")
heading(doc, "Practice for stronger schools")
body(doc, "FHTA promotes leadership approaches aligned with the changing needs of learners, teachers and communities.", True)
for title, detail in [
    ("Instructional leadership", "Keep teaching quality and learning outcomes at the centre."),
    ("Digital integration", "Use appropriate technology to improve access, administration and learning."),
    ("Inclusive education", "Advance equity, participation and support for every learner."),
    ("Community engagement", "Build productive relationships with parents and stakeholders."),
    ("Data-informed decisions", "Use evidence for planning, accountability and improvement."),
    ("Student wellbeing", "Support learners' safety, belonging and holistic development."),
]:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, PALE)
    margins(cell, 115, 150, 115, 150)
    p = cell.paragraphs[0]
    add_text(p, title, 10, NAVY, True)
    p2 = cell.add_paragraph()
    add_text(p2, detail, 8.8, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Page 7
doc.add_page_break()
kicker(doc, "Opportunity for partnership")
heading(doc, "A practical bridge between policy and schools")
body(doc, "FHTA is positioned to support effective dialogue and implementation by connecting national priorities with the experience of primary school leaders.", True)
heading(doc, "Partnership opportunities", 2)
bullet(doc, "Structured consultation with Heads of Schools before and during major policy initiatives.")
bullet(doc, "Joint professional learning on leadership, curriculum, inclusion, wellbeing and digital transformation.")
bullet(doc, "District-level feedback loops that identify implementation barriers and promising practice early.")
bullet(doc, "Coordinated communication of key notices, resources and opportunities through the FHTA network and portal.")
bullet(doc, "Collaborative recognition of school leadership excellence and innovation.")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "A shared goal", 10, BLUE, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "Confident school leaders, supported teachers,\nand better outcomes for Fiji's children.", 17, NAVY, True)

# Back cover
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.add_run().add_picture(str(LOGO), width=Inches(1.45))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
add_text(p, "FIJI HEAD TEACHERS ASSOCIATION", 11, BLUE, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
add_text(p, "A national network for primary\nschool leadership", 21, NAVY, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
add_text(p, "Connect with FHTA", 12, NAVY, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p, "fijiheadteachersassociation@gmail.com", 10, BLUE, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
add_text(p, "Website services include news, events, resources,\nmember support, representatives and professional learning.", 9.5, INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(26)
add_text(p, "Prepared from the current FHTA website and platform capabilities.", 7.5, "6B7C8F", False, True)

doc.core_properties.title = "FHTA Minister for Education Briefing Booklet"
doc.core_properties.subject = "Fiji Head Teachers Association website and partnership briefing"
doc.core_properties.author = "Fiji Head Teachers Association"
path = OUT / "FHTA_Minister_Briefing_Booklet.docx"
doc.save(path)
print(path)
